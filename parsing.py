"""
Resume document parsing: PDF (PyMuPDF) and DOCX (python-docx).

Shared by the Celery worker and the Streamlit local-fallback mode.
"""

import io
import re
import os

SUPPORTED_EXTENSIONS = (".pdf", ".docx")

# PostgreSQL TEXT columns reject NUL (0x00) outright, and psycopg2 cannot
# encode lone UTF-16 surrogates. Both show up routinely in real-world PDFs
# (broken CID/ToUnicode maps, embedded form fields, scanner artefacts), so
# every parsed string is scrubbed before it can reach the database.
_CONTROL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def sanitize_text(text: str) -> str:
    """Make parsed document text safe to persist and to JSON-encode.

    Strips NUL/control characters and lone surrogates, and normalizes
    whitespace runs. Never raises - a resume must not 500 the API.
    """
    if not text:
        return ""
    # Drop unpaired surrogates (psycopg2/JSON both choke on them).
    text = text.encode("utf-8", "ignore").decode("utf-8", "ignore")
    text = _CONTROL_RE.sub(" ", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t\u00a0]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def parse_resume(filename: str, payload: bytes) -> str:
    """Extract plain text from a resume given its filename and raw bytes.

    Raises ValueError for unsupported types or unparseable content.
    """
    ext = os.path.splitext(filename or "")[1].lower()

    if ext == ".pdf":
        return sanitize_text(_parse_pdf(payload))
    if ext == ".docx":
        return sanitize_text(_parse_docx(payload))
    raise ValueError(
        f"Unsupported file type '{ext or 'unknown'}'. "
        f"Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
    )


def _parse_pdf(payload: bytes) -> str:
    try:
        import pymupdf as fitz   # modern module name (PyMuPDF >= 1.24);
    except ImportError:          # pragma: no cover - older builds
        import fitz              # legacy shim (deprecation-warns on 1.28)
    doc = fitz.open(stream=payload, filetype="pdf")
    try:
        # Newline-joined: extract_name() and the education heuristics are
        # line-oriented, and " ".join() used to collapse the whole CV onto a
        # single line, which broke name extraction for every PDF resume.
        return "\n".join(page.get_text() for page in doc)
    finally:
        doc.close()


def _parse_docx(payload: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise ValueError(
            "DOCX support requires python-docx to be installed"
        ) from exc

    document = Document(io.BytesIO(payload))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    # Include table cell text — resumes love tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)